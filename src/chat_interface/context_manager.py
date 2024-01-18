import os
from typing import List, Dict
from docx import Document
from utils.summarizer import summarize_conversation


class ContextManager:
    def __init__(self) -> None:
        self.raw_context: List[Dict[str, str]] = []
        self.summarized_context: str = ""
        self.max_recent_messages: int = 3

    def update_context(self, user_message: str, assistant_message: str) -> None:
        self.raw_context.append({"role": "user", "content": user_message})
        self.raw_context.append(
            {"role": "assistant", "content": assistant_message})

        if len(self.raw_context) > self.max_recent_messages * 2:
            full_conversation: str = ' '.join(
                [msg['content'] for msg in self.raw_context])
            self.summarized_context = summarize_conversation(full_conversation)
            self.raw_context = self.raw_context[-self.max_recent_messages * 2:]

    def get_context(self) -> List[str]:
        if self.summarized_context:
            return [self.summarized_context] + [msg['content'] for msg in self.raw_context]
        else:
            return [msg['content'] for msg in self.raw_context]
        
    def get_context_api(self) -> List[Dict[str, str]]:
        formatted_context = []
        if self.summarized_context:
            formatted_context.append(
                {"role": "system", "content": self.summarized_context})
        formatted_context.extend(self.raw_context)
        return formatted_context

    def reduce_context(self) -> None:
        if len(self.raw_context) > 2:
            self.raw_context = self.raw_context[-2:]

    def save_conversation_as_docx(self, summary: str, details: List[str], file_name: str) -> None:
        document = Document()
        document.add_heading('Conversation Summary', level=1)
        document.add_paragraph(summary)
        document.add_heading('Conversation Details', level=2)
        for detail in details:
            document.add_paragraph(detail)

        desktop_path = os.path.join(os.path.join(
            os.environ['USERPROFILE']), 'Desktop')
        file_path = os.path.join(
            desktop_path, 'chat-ai', 'conversations', f'{file_name}.docx')
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        document.save(file_path)

    def reset_context(self) -> None:
        self.raw_context = []
        self.summarized_context = ""
